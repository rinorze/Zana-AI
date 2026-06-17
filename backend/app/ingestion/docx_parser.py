"""Extract plain text from uploaded DOCX files."""
from __future__ import annotations

from io import BytesIO

from docx import Document


def parse_docx(content: bytes) -> str:
    if not content:
        return ""
    doc = Document(BytesIO(content))

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)
